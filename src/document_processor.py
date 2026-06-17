"""
document_processor.py
Loads PDF / TXT / DOCX / CSV files and splits them into LangChain Documents.
"""
from __future__ import annotations
import csv
import io
from pathlib import Path
from typing import List

from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter


class DocumentProcessor:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 64):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
            length_function=len,
        )

    # ── Public ────────────────────────────────────────────────────────────────
    def load_and_split(self, file_path: str, source_name: str = "") -> List[Document]:
        """Return a list of chunked Documents from any supported file type."""
        path = Path(file_path)
        ext = path.suffix.lower()
        source_name = source_name or path.name

        loaders = {
            ".pdf": self._load_pdf,
            ".txt": self._load_txt,
            ".docx": self._load_docx,
            ".csv": self._load_csv,
        }
        loader_fn = loaders.get(ext)
        if loader_fn is None:
            raise ValueError(f"Unsupported file type: {ext}")

        raw_docs = loader_fn(file_path, source_name)
        chunks = self.splitter.split_documents(raw_docs)

        # Tag every chunk with its source
        for i, chunk in enumerate(chunks):
            chunk.metadata.setdefault("source", source_name)
            chunk.metadata["chunk_index"] = i

        return chunks

    # ── Loaders ───────────────────────────────────────────────────────────────
    def _load_pdf(self, path: str, source: str) -> List[Document]:
        """PyMuPDF loader (fallback to pdfplumber)."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            pages = []
            for i, page in enumerate(doc):
                text = page.get_text("text").strip()
                if text:
                    pages.append(Document(
                        page_content=text,
                        metadata={"source": source, "page": i + 1},
                    ))
            doc.close()
            return pages
        except ImportError:
            pass

        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = (page.extract_text() or "").strip()
                    if text:
                        pages.append(Document(
                            page_content=text,
                            metadata={"source": source, "page": i + 1},
                        ))
            return pages
        except ImportError:
            raise ImportError("Install PyMuPDF (`pip install pymupdf`) or pdfplumber (`pip install pdfplumber`) to read PDFs.")

    def _load_txt(self, path: str, source: str) -> List[Document]:
        with open(path, encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return [Document(page_content=text, metadata={"source": source})]

    def _load_docx(self, path: str, source: str) -> List[Document]:
        try:
            import docx
        except ImportError:
            raise ImportError("Install python-docx: `pip install python-docx`")
        doc = docx.Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        return [Document(page_content=text, metadata={"source": source})]

    def _load_csv(self, path: str, source: str) -> List[Document]:
        docs = []
        with open(path, encoding="utf-8", errors="ignore", newline="") as f:
            reader = csv.DictReader(f)
            for i, row in enumerate(reader):
                text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                if text:
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": source, "row": i + 1},
                    ))
        return docs
